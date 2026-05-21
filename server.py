from __future__ import annotations

import shutil
import tempfile
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pdf2docx import Converter

app = FastAPI(title="ToolKit Pro PDF to Word Engine", version="1.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


def _merge_page_docs(page_docs: list[Path], out_path: Path) -> None:
    from docxcompose.composer import Composer

    master = Document(str(page_docs[0]))
    composer = Composer(master)
    for doc_path in page_docs[1:]:
        composer.append(Document(str(doc_path)))
    composer.save(str(out_path))


def _fallback_page_image_doc(pdf_path: Path, page_number: int, out_docx: Path) -> None:
    pdf = fitz.open(str(pdf_path))
    page = pdf[page_number]
    pix = page.get_pixmap(dpi=220, alpha=False)
    img_path = out_docx.with_suffix('.png')
    pix.save(str(img_path))

    doc = Document()
    section = doc.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin
    doc.add_picture(str(img_path), width=usable_width)
    doc.save(str(out_docx))


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/api/pdf-to-word")
async def pdf_to_word(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file name provided.")

    source_name = Path(file.filename).name
    if not source_name.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported.")

    tmp_dir = Path(tempfile.mkdtemp(prefix="toolkitpro_pdf2docx_"))
    in_path = tmp_dir / source_name
    out_path = tmp_dir / f"{Path(source_name).stem}_editable.docx"

    try:
        with in_path.open("wb") as f:
            shutil.copyfileobj(file.file, f)

        # Convert page-by-page to avoid random page loss on complex PDFs.
        pdf = fitz.open(str(in_path))
        page_count = pdf.page_count
        pdf.close()

        page_docs: list[Path] = []
        failed_pages: list[int] = []

        for i in range(page_count):
            page_doc = tmp_dir / f"page_{i+1:04d}.docx"
            try:
                converter = Converter(str(in_path))
                try:
                    converter.convert(str(page_doc), start=i, end=i + 1)
                finally:
                    converter.close()

                if not page_doc.exists() or page_doc.stat().st_size == 0:
                    raise RuntimeError("empty docx page output")
            except Exception:
                failed_pages.append(i + 1)
                _fallback_page_image_doc(in_path, i, page_doc)

            page_docs.append(page_doc)

        if not page_docs:
            raise HTTPException(status_code=500, detail="Conversion failed: no output pages generated.")

        _merge_page_docs(page_docs, out_path)

        if not out_path.exists() or out_path.stat().st_size == 0:
            raise HTTPException(status_code=500, detail="Conversion failed: output document is empty.")

        docx_bytes = out_path.read_bytes()
        headers = {
            "Content-Disposition": f"attachment; filename=\"{out_path.name}\"",
            "X-Conversion-Failed-Pages": ",".join(map(str, failed_pages)) if failed_pages else "",
        }
        return StreamingResponse(
            iter([docx_bytes]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers,
        )
    except HTTPException:
        raise
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=500, detail=f"Conversion error: {exc}") from exc
    finally:
        try:
            await file.close()
        except Exception:
            pass
        shutil.rmtree(tmp_dir, ignore_errors=True)


if __name__ == "__main__":
    import uvicorn

    uvicorn.run("server:app", host="127.0.0.1", port=8000, reload=False)
